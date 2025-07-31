import fitz
import os
import re
import hashlib
import time
from datetime import datetime
from typing import List, Dict, Any, Optional
from langchain .text_splitter import RecursiveCharacterTextSplitter
from langchain .schema import Document
import streamlit as st
from enum import Enum
from dataclasses import dataclass

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ProcessingStage (Enum):
    INITIALIZING = "initializing"
    READING_PDF = "reading_pdf"
    READING_DOCX = "reading_docx"
    READING_TXT = "reading_txt"
    PROCESSING_TEXT = "processing_text"
    CREATING_CHUNKS = "creating_chunks"
    GENERATING_EMBEDDINGS = "generating_embeddings"
    STORING_DOCUMENTS = "storing_documents"
    COMPLETED = "completed"
    ERROR = "error"


@dataclass
class ProgressState:
    stage: ProcessingStage = ProcessingStage .INITIALIZING
    current_file: str = ""
    current_step: int = 0
    total_steps: int = 0
    current_file_index: int = 0
    total_files: int = 0
    message: str = ""
    error_message: str = ""
    start_time: float = 0
    stage_start_time: float = 0

    @property
    def progress_percent(self) -> float:
        if self .total_steps == 0:
            return 0.0
        return min(100.0, (self .current_step / self .total_steps)*100.0)

    @property
    def file_progress_percent(self) -> float:
        if self .total_files == 0:
            return 0.0
        return min(100.0, (self .current_file_index / self .total_files)*100.0)

    @property
    def elapsed_time(self) -> float:
        return time .time()-self .start_time if self .start_time > 0 else 0

    @property
    def stage_elapsed_time(self) -> float:
        return time .time()-self .stage_start_time if self .stage_start_time > 0 else 0


class SimpleProgressTracker:
    """Упрощенный трекер прогресса для v1.5.0, интегрированный со Streamlit"""

    def __init__(self):
        self .state = ProgressState()
        self .progress_placeholder = None
        self .status_placeholder = None
        self .detail_placeholder = None
        self ._is_active = False

    def setup_ui(self, container=None):
        """Настройка UI элементов для отображения прогресса"""
        if container is None:
            container = st .container()

        with container:
            self .progress_placeholder = st .empty()
            self .status_placeholder = st .empty()
            self .detail_placeholder = st .empty()

    def start_session(self, total_files: int, title: str = "Обработка документов"):
        """Начать сессию обработки"""
        self .state = ProgressState(
            total_files=total_files,
            start_time=time .time(),
            stage_start_time=time .time()
        )
        self ._is_active = True

        if self .progress_placeholder is None:
            self .setup_ui()

        self ._update_display(f"{title} ({total_files} файлов)")

    def start_file(self, filename: str, total_steps: int = 5):
        """Начать обработку файла"""
        self .state .current_file = filename
        self .state .current_file_index += 1
        self .state .current_step = 0
        self .state .total_steps = total_steps
        self .state .stage = ProcessingStage .READING_PDF
        self .state .stage_start_time = time .time()
        self .state .message = f"Начинаю обработку файла {filename}"
        self .state .error_message = ""

        self ._update_display()

    def update_stage(self, stage: ProcessingStage, message: str = "", step_increment: int = 1):
        """Обновить текущий этап"""
        self .state .stage = stage
        self .state .current_step += step_increment
        self .state .stage_start_time = time .time()

        if message:
            self .state .message = message
        else:
            self .state .message = self ._get_default_message(stage)

        self ._update_display()

    def update_progress(self, current: int, total: int = None, message: str = ""):
        """Обновить прогресс текущего этапа"""
        self .state .current_step = current
        if total is not None:
            self .state .total_steps = total
        if message:
            self .state .message = message

        self ._update_display()

    def set_error(self, error_message: str):
        """Зарегистрировать ошибку"""
        self .state .stage = ProcessingStage .ERROR
        self .state .error_message = error_message
        self .state .message = f"Ошибка: {error_message}"

        self ._update_display()

    def complete_file(self):
        """Завершить обработку текущего файла"""
        self .state .stage = ProcessingStage .COMPLETED
        self .state .current_step = self .state .total_steps
        self .state .message = f"Файл {self .state .current_file} обработан успешно"

        self ._update_display()

    def finish_session(self, success: bool = True):
        """Завершить сессию"""
        self ._is_active = False

        if success:
            self ._show_final_success()
        else:
            self ._show_final_error()

    def _update_display(self, title: str = None):
        """Обновить отображение прогресса"""
        if not self ._is_active or not self .progress_placeholder:
            return

        try:
            with self .progress_placeholder:
                if title:
                    st .subheader(title)

                if self .state .total_files > 0:
                    file_progress = self .state .file_progress_percent / 100
                    st .progress(
                        file_progress,
                        text=f"Файл {self .state .current_file_index}/{self .state .total_files}"
                    )

            with self .status_placeholder:
                if self .state .current_file:
                    step_progress = self .state .progress_percent / 100

                    st .progress(
                        step_progress,
                        text=f"{self .state .current_file}: {self .state .message} ({self .state .progress_percent:.0f}%)"
                    )

            with self .detail_placeholder:
                if self .state .error_message:
                    st .error(f"Ошибка: {self .state .error_message}")
                elif self .state .stage_elapsed_time > 0:
                    elapsed = self .state .stage_elapsed_time
                    if elapsed < 60:
                        time_str = f"{elapsed:.1f} сек"
                    else:
                        time_str = f"{elapsed / 60:.1f} мин"

                    st .caption(f"Время этапа: {time_str}")

        except Exception as e:
            pass

    def _show_final_success(self):
        """Показать финальное сообщение об успехе"""
        if self .progress_placeholder:
            with self .progress_placeholder:
                st .success(
                    f"Успешно обработано {self .state .total_files} файлов")

        if self .status_placeholder:
            with self .status_placeholder:
                total_time = self .state .elapsed_time
                if total_time < 60:
                    time_str = f"{total_time:.1f} секунд"
                else:
                    time_str = f"{total_time / 60:.1f} минут"
                st .info(f"Общее время обработки: {time_str}")

        if self .detail_placeholder:
            with self .detail_placeholder:
                st .empty()

    def _show_final_error(self):
        """Показать финальное сообщение об ошибке"""
        if self .progress_placeholder:
            with self .progress_placeholder:
                st .error("Обработка завершена с ошибками")

    def _get_default_message(self, stage: ProcessingStage) -> str:
        """Получить сообщение по умолчанию для этапа"""
        message_map = {
            ProcessingStage .INITIALIZING: "Инициализация",
            ProcessingStage .READING_PDF: "Чтение PDF файла",
            ProcessingStage .READING_DOCX: "Чтение DOCX файла",
            ProcessingStage .READING_TXT: "Чтение TXT файла",
            ProcessingStage .PROCESSING_TEXT: "Обработка текста",
            ProcessingStage .CREATING_CHUNKS: "Создание фрагментов",
            ProcessingStage .GENERATING_EMBEDDINGS: "Генерация эмбеддингов",
            ProcessingStage .STORING_DOCUMENTS: "Сохранение в базу данных",
            ProcessingStage .COMPLETED: "Завершено",
            ProcessingStage .ERROR: "Ошибка"
        }
        return message_map .get(stage, "Обработка")

    def get_current_state(self) -> Dict[str, Any]:
        """Получить текущее состояние для отладки"""
        return {
            "stage": self .state .stage .value,
            "current_file": self .state .current_file,
            "progress_percent": self .state .progress_percent,
            "file_progress_percent": self .state .file_progress_percent,
            "current_step": self .state .current_step,
            "total_steps": self .state .total_steps,
            "current_file_index": self .state .current_file_index,
            "total_files": self .state .total_files,
            "message": self .state .message,
            "error_message": self .state .error_message,
            "elapsed_time": self .state .elapsed_time,
            "is_active": self ._is_active
        }


class ProgressContext:
    """Контекстный менеджер для автоматического управления прогрессом"""

    def __init__(self, total_files: int, title: str = "Обработка документов"):
        self .tracker = SimpleProgressTracker()
        self .total_files = total_files
        self .title = title
        self .success = True

    def __enter__(self):
        self .tracker .setup_ui()
        self .tracker .start_session(self .total_files, self .title)
        return self .tracker

    def __exit__(self, exc_type, exc_val, exc_tb):
        if exc_type is not None:
            self .success = False
            if hasattr(exc_val, '__str__'):
                self .tracker .set_error(str(exc_val))

        self .tracker .finish_session(self .success)
        return False


class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200, progress_tracker: Optional[SimpleProgressTracker] = None):
        self .chunk_size = chunk_size
        self .chunk_overlap = chunk_overlap
        self .progress_tracker = progress_tracker
        self .text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\nСтатья", "\n\n", ".\n", "\n", ".", "", ""],
            keep_separator=True
        )

    def update_chunk_settings(self, chunk_size: int, chunk_overlap: int):
        """Обновить настройки разбиения на фрагменты"""
        self .chunk_size = chunk_size
        self .chunk_overlap = chunk_overlap
        self .text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\nСтатья", "\n\n", ".\n", "\n", ".", "", ""],
            keep_separator=True
        )

    def extract_text_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        try:
            if self .progress_tracker:
                self .progress_tracker .update_stage(
                    ProcessingStage .READING_PDF, "Открытие PDF файла")

            doc = fitz .open(pdf_path)
            text = ""
            metadata = {
                "filename": os .path .basename(pdf_path),
                "page_count": len(doc),
                "file_path": pdf_path
            }

            if self .progress_tracker:
                self .progress_tracker .update_stage(
                    ProcessingStage .PROCESSING_TEXT, f"Извлечение текста из {len(doc)} страниц")

            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page .get_text().strip()
                if page_text:
                    text += f"\n\n{page_text}"

                if self .progress_tracker:
                    progress = int(((page_num + 1)/len(doc))*100)
                    self .progress_tracker .update_progress(
                        page_num + 1, len(doc), f"Обработана страница {page_num + 1} из {len(doc)}")

            doc .close()

            return {
                "text": text .strip(),
                "metadata": metadata
            }

        except Exception as e:
            error_msg = f"Error extracting text from {pdf_path}: {str(e)}"
            if self .progress_tracker:
                self .progress_tracker .set_error(error_msg)
            st .error(error_msg)
            return None

    def extract_text_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Извлечение текста из DOCX файла"""
        try:
            if not DOCX_AVAILABLE:
                error_msg = "python-docx library not installed. Please install: pip install python-docx"
                if self .progress_tracker:
                    self .progress_tracker .set_error(error_msg)
                else:
                    st .error(error_msg)
                return None

            if self .progress_tracker:
                self .progress_tracker .update_stage(
                    ProcessingStage .READING_DOCX, "Открытие DOCX файла")

            doc = DocxDocument(docx_path)
            text = ""

            if self .progress_tracker:
                total_paragraphs = len(doc .paragraphs)
                self .progress_tracker .update_stage(
                    ProcessingStage .PROCESSING_TEXT, f"Извлечение текста из {total_paragraphs} параграфов")

            for i, paragraph in enumerate(doc .paragraphs):
                if paragraph .text .strip():
                    text += paragraph .text + "\n\n"

                if self .progress_tracker:
                    progress = int(((i + 1)/len(doc .paragraphs))*100)
                    self .progress_tracker .update_progress(
                        i + 1, len(doc .paragraphs), f"Обработан параграф {i + 1} из {len(doc .paragraphs)}")

            for table in doc .tables:
                for row in table .rows:
                    row_text = []
                    for cell in row .cells:
                        if cell .text .strip():
                            row_text .append(cell .text .strip())
                    if row_text:
                        text += " |".join(row_text)+"\n"

            metadata = {
                "filename": os .path .basename(docx_path),
                "page_count": "N/A",
                "file_path": docx_path,
                "document_type": "docx",
                "paragraphs_count": len(doc .paragraphs),
                "tables_count": len(doc .tables)
            }

            return {
                "text": text .strip(),
                "metadata": metadata
            }

        except Exception as e:
            error_msg = f"Error extracting text from {docx_path}: {str(e)}"
            if self .progress_tracker:
                self .progress_tracker .set_error(error_msg)
            else:
                st .error(error_msg)
            return None

    def extract_text_from_txt(self, txt_path: str) -> Dict[str, Any]:
        """Извлечение текста из TXT файла"""
        try:
            if self .progress_tracker:
                self .progress_tracker .update_stage(
                    ProcessingStage .READING_TXT, "Открытие TXT файла")

            encodings = ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']
            text = ""
            used_encoding = None

            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding)as file:
                        text = file .read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue

            if not text:
                error_msg = f"Could not decode file {txt_path} with any of the attempted encodings"
                if self .progress_tracker:
                    self .progress_tracker .set_error(error_msg)
                else:
                    st .error(error_msg)
                return None

            if self .progress_tracker:
                lines_count = len(text .split('\n'))
                chars_count = len(text)
                self .progress_tracker .update_stage(
                    ProcessingStage .PROCESSING_TEXT, f"Обработка текста ({lines_count} строк, {chars_count} символов)")

            file_stats = os .stat(txt_path)

            metadata = {
                "filename": os .path .basename(txt_path),
                "page_count": "N/A",
                "file_path": txt_path,
                "document_type": "txt",
                "encoding": used_encoding,
                "lines_count": len(text .split('\n')),
                "characters_count": len(text),
                "size_bytes": file_stats .st_size
            }

            return {
                "text": text .strip(),
                "metadata": metadata
            }

        except Exception as e:
            error_msg = f"Error extracting text from {txt_path}: {str(e)}"
            if self .progress_tracker:
                self .progress_tracker .set_error(error_msg)
            else:
                st .error(error_msg)
            return None

    def chunk_document(self, text: str, metadata: Dict[str, Any]) -> List[Document]:
        if self .progress_tracker:
            self .progress_tracker .update_stage(
                ProcessingStage .CREATING_CHUNKS, "Разбиение текста на фрагменты")

        chunks = self .text_splitter .split_text(text)
        valid_chunks = [c for c in chunks if len(c .strip()) >= 50]

        if self .progress_tracker:
            self .progress_tracker .update_progress(
                0, len(valid_chunks), f"Создание {len(valid_chunks)} фрагментов")

        documents = []
        valid_chunk_id = 0

        for i, chunk in enumerate(chunks):
            chunk_cleaned = chunk .strip()
            if len(chunk_cleaned) >= 50:
                chunk_metadata = metadata .copy()
                chunk_metadata["chunk_id"] = valid_chunk_id
                chunk_metadata["chunk_size"] = len(chunk_cleaned)

                documents .append(Document(
                    page_content=chunk_cleaned,
                    metadata=chunk_metadata
                ))
                valid_chunk_id += 1

                if self .progress_tracker:
                    self .progress_tracker .update_progress(valid_chunk_id, len(
                        valid_chunks), f"Создан фрагмент {valid_chunk_id} из {len(valid_chunks)}")

        return documents

    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """Универсальный метод для извлечения текста из любого поддерживаемого формата"""
        file_extension = os .path .splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            return self .extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self .extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self .extract_text_from_txt(file_path)
        else:
            error_msg = f"Unsupported file format: {file_extension}"
            if self .progress_tracker:
                self .progress_tracker .set_error(error_msg)
            else:
                st .error(error_msg)
            return None

    def process_file(self, file_path: str) -> List[Document]:
        """Универсальный метод для обработки файла любого поддерживаемого формата"""
        extracted_data = self .extract_text_from_file(file_path)
        if not extracted_data:
            return []

        documents = self .chunk_document(
            extracted_data["text"],
            extracted_data["metadata"]
        )

        return documents

    def process_pdf_file(self, pdf_path: str) -> List[Document]:
        extracted_data = self .extract_text_from_pdf(pdf_path)
        if not extracted_data:
            return []

        documents = self .chunk_document(
            extracted_data["text"],
            extracted_data["metadata"]
        )

        return documents

    def process_directory(self, directory_path: str) -> List[Document]:
        """Обработка всех поддерживаемых файлов в директории"""
        if not os .path .exists(directory_path):
            st .error(f"Directory {directory_path} does not exist")
            return []

        supported_extensions = ['.pdf', '.docx', '.txt']
        supported_files = [
            f for f in os .listdir(directory_path)
            if any(f .lower().endswith(ext)for ext in supported_extensions)
        ]

        if not supported_files:
            st .warning(
                f"No supported files (PDF, DOCX, TXT) found in {directory_path}")
            return []

        all_documents = []

        if self .progress_tracker:
            if not self .progress_tracker ._is_active:
                self .progress_tracker .start_session(
                    len(supported_files), f"Обработка {len(supported_files)} файлов")

            for i, file_name in enumerate(supported_files):
                try:
                    self .progress_tracker .start_file(file_name, 5)
                    file_path = os .path .join(directory_path, file_name)

                    documents = self .process_file(file_path)
                    all_documents .extend(documents)

                    self .progress_tracker .complete_file()

                except Exception as e:
                    self .progress_tracker .set_error(
                        f"Ошибка обработки {file_name}: {str(e)}")
                    continue

        else:
            progress_bar = st .progress(0)
            status_text = st .empty()

            for i, file_name in enumerate(supported_files):
                status_text .text(f"Processing {file_name}...")
                file_path = os .path .join(directory_path, file_name)

                documents = self .process_file(file_path)
                all_documents .extend(documents)

                progress_bar .progress((i + 1)/len(supported_files))

            status_text .text(
                f"Processed {len(supported_files)} files, created {len(all_documents)} chunks")

        return all_documents

    def process_pdf_directory(self, directory_path: str) -> List[Document]:
        """Backward compatibility method - теперь использует универсальный process_directory"""
        return self .process_directory(directory_path)

    def _sanitize_filename(self, filename: str) -> str:
        """Очищает имя файла от недопустимых символов"""
        name_without_ext = os .path .splitext(filename)[0]

        clean_name = re .sub(r'[^\w\s\-\.]', '', name_without_ext)
        clean_name = re .sub(r'\s+', '_', clean_name)
        clean_name = re .sub(r'-+', '_', clean_name)
        clean_name = re .sub(r'_+', '_', clean_name)

        clean_name = clean_name .strip('_')

        if len(clean_name) > 50:
            clean_name = clean_name[:50]

        if len(clean_name) < 1:
            clean_name = f"document_{datetime .now().strftime('%Y%m%d_%H%M%S')}"

        return clean_name

    def _get_unique_filepath(self, directory: str, filename: str) -> str:
        base_path = os .path .join(directory, filename)

        if not os .path .exists(base_path):
            return base_path

        name_without_ext = os .path .splitext(filename)[0]
        extension = os .path .splitext(filename)[1]

        counter = 1
        while True:
            new_filename = f"{name_without_ext}_({counter}){extension}"
            new_path = os .path .join(directory, new_filename)
            if not os .path .exists(new_path):
                return new_path
            counter += 1

    def process_uploaded_file(self, uploaded_file) -> List[Document]:
        try:
            if self .progress_tracker:
                self .progress_tracker .start_file(uploaded_file .name, 5)

            docs_dir = os .path .abspath("./data/documents")
            os .makedirs(docs_dir, exist_ok=True)

            file_content = uploaded_file .getbuffer()

            sanitized_name = self ._sanitize_filename(uploaded_file .name)
            file_extension = os .path .splitext(uploaded_file .name)[1].lower()
            final_filename = sanitized_name + file_extension

            final_path = self ._get_unique_filepath(docs_dir, final_filename)
            final_filename = os .path .basename(final_path)

            with open(final_path, "wb")as f:
                f .write(file_content)

            if not self .progress_tracker:
                st .info(f"Файл сохранен: {final_filename}")

            extracted_data = self .extract_text_from_file(final_path)
            if not extracted_data:
                file_type = os .path .splitext(final_path)[1].upper()
                error_msg = f"Не удалось извлечь текст из {file_type} файла"
                if self .progress_tracker:
                    self .progress_tracker .set_error(error_msg)
                else:
                    st .error(error_msg)
                return []

            extracted_data["metadata"]["filename"] = final_filename
            extracted_data["metadata"]["file_path"] = final_path
            extracted_data["metadata"]["original_name"] = uploaded_file .name

            documents = self .chunk_document(
                extracted_data["text"],
                extracted_data["metadata"]
            )

            if self .progress_tracker:
                self .progress_tracker .complete_file()
            elif documents:
                page_info = extracted_data['metadata'].get('page_count', 'N/A')
                if page_info != 'N/A':
                    st .success(
                        f"Создано {len(documents)} фрагментов из {page_info} страниц")
                else:
                    st .success(f"Создано {len(documents)} фрагментов")

            return documents

        except Exception as e:
            error_msg = f"Ошибка обработки файла: {str(e)}"
            if self .progress_tracker:
                self .progress_tracker .set_error(error_msg)
            else:
                st .error(error_msg)
            return []

    def get_document_summary(self, documents: List[Document]) -> Dict[str, Any]:
        if not documents:
            return {"total_chunks": 0, "total_characters": 0, "unique_files": 0}

        total_chunks = len(documents)
        total_characters = sum(len(doc .page_content)for doc in documents)
        unique_files = len(set(doc .metadata .get("filename", "")
                               for doc in documents))

        return {
            "total_chunks": total_chunks,
            "total_characters": total_characters,
            "unique_files": unique_files,
            "average_chunk_size": total_characters / total_chunks if total_chunks > 0 else 0
        }

    def delete_physical_file(self, filename: str) -> bool:
        try:
            file_path = os .path .join("./data/documents", filename)

            if os .path .exists(file_path):
                os .remove(file_path)
                st .success(f"Physical file deleted: {filename}")
                return True
            else:
                st .warning(f"Physical file not found: {filename}")
                return False

        except Exception as e:
            st .error(f"Error deleting physical file {filename}: {str(e)}")
            return False

    def rename_physical_file(self, old_filename: str, new_filename: str) -> bool:
        try:
            docs_dir = "./data/documents"
            old_path = os .path .join(docs_dir, old_filename)
            new_path = os .path .join(docs_dir, new_filename)

            if not os .path .exists(old_path):
                st .warning(f"Original file not found: {old_filename}")
                return False

            if os .path .exists(new_path):
                st .error(f"File with name {new_filename} already exists")
                return False

            os .rename(old_path, new_path)
            st .success(f"File renamed: {old_filename} → {new_filename}")
            return True

        except Exception as e:
            st .error(f"Error renaming file: {str(e)}")
            return False

    def get_physical_file_info(self, filename: str) -> Dict[str, Any]:
        try:
            file_path = os .path .join("./data/documents", filename)

            if not os .path .exists(file_path):
                return {"exists": False, "path": file_path}

            stat_info = os .stat(file_path)

            return {
                "exists": True,
                "path": file_path,
                "size_bytes": stat_info .st_size,
                "size_mb": round(stat_info .st_size / (1024 * 1024), 2),
                "modified_time": datetime .fromtimestamp(stat_info .st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            }

        except Exception as e:
            st .error(f"Error getting file info for {filename}: {str(e)}")
            return {"exists": False, "error": str(e)}
